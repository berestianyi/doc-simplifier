import os
import re
from typing import List, Dict

from config import settings

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from python_docx_replace import docx_replace

from contracts.services.dto import VehicleData


class DocxEditor:
    font_cache = None
    vehicle_table_index = None

    def __init__(self, template_path: str):
        self.doc = Document(template_path)

    def _add_vehicle_row(self, table, index: int, vehicle: VehicleData):
        row = table.add_row().cells
        brand_model = f"{vehicle['brand']}, {vehicle['model']}"
        self._set_cell_text(row[0], str(index))
        self._set_cell_text(row[1], brand_model)
        self._set_cell_text(row[2], vehicle["number"])
        self._set_cell_text(row[3], str(vehicle["year"]))

    def _set_cell_text(self, cell, text: str):
        paragraph = cell.paragraphs[0]
        paragraph.clear()
        run = paragraph.add_run(text)
        run.bold = True
        if self.font_cache:
            run.font.name = self.font_cache

    def _apply_table_styles(self, table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        font = run.font
                        font.name = self.font_cache

    def replace_text(self, replacements: Dict[str, str]):
        docx_replace(self.doc, **replacements)

    def add_table(self, data: List[VehicleData]):
        table = self.doc.tables[self.vehicle_table_index]
        for idx, vehicle in enumerate(data, start=1):
            self._add_vehicle_row(table, idx, vehicle)
        self._apply_table_styles(table)

    def save(self, filename: str) -> str:
        output_path = os.path.join(settings.MEDIA_ROOT, "documents", self._sanitize_filename(filename))
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
        return output_path

    @staticmethod
    def _sanitize_filename(filename: str) -> str:
        return re.sub(r'[<>:"/\\|?*]', '_', filename).strip()


class RoyalDocxEditor(DocxEditor):
    font_cache = 'Times New Roman'
    vehicle_table_index = 1


class RolandDocxEditor(DocxEditor):
    font_cache = 'Arial'
    vehicle_table_index = 0
