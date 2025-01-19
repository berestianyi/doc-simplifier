import os

import pymorphy3

from django.db.models import Q
from django.forms import model_to_dict
from django.http import Http404, FileResponse, HttpResponse
from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse_lazy, reverse
from django.utils.functional import cached_property
from django.views.generic import ListView, View, CreateView, UpdateView, TemplateView

from docx import Document
from python_docx_replace import docx_replace

from business_entities.models import BusinessEntities
from business_entities.views import BusinessEntityMixin
from config import settings
from documents.models import Documents
from vehicles.models import Vehicles
from vehicles.views import VehicleMixin
from .forms import ContractTimeRangeForm, TemplatesForm
from .models import Templates, Contracts

morph = pymorphy3.MorphAnalyzer(lang='uk')


class TemplateMixin:
    @cached_property
    def template_obj(self):
        template_id = self.kwargs.get("template_id")
        return get_object_or_404(Templates, pk=template_id)


class ReplacementManager:
    MONTH_NAMES_UA = {
        1: "січня",
        2: "лютого",
        3: "березня",
        4: "квітня",
        5: "травня",
        6: "червня",
        7: "липня",
        8: "серпня",
        9: "вересня",
        10: "жовтня",
        11: "листопада",
        12: "грудня",
    }

    def __init__(
            self, start_date, expire_date, business_entity, vehicles, bank
    ):
        self.start_date = start_date
        self.expire_date = expire_date
        self.bank = bank
        self.business_entity = business_entity
        self.vehicles = vehicles

    @classmethod
    def _model_to_dict_or_empty(cls, instance, fields, default="") -> dict:
        if instance is not None:
            return model_to_dict(instance, fields=fields)
        return {field: default for field in fields}

    @classmethod
    def upper_or_empty_str(cls, text: str) -> str:
        return text.upper() if text else ""

    @classmethod
    def title_or_empty_str(cls, text: str) -> str:
        return text.title() if text else ""

    @classmethod
    def to_genitive(cls, full_name: str) -> str:
        tokens = full_name.split()
        result_tokens = []

        for token in tokens:
            parsed = morph.parse(token)
            best_parse = parsed[0]
            gent_form = best_parse.inflect({'gent'})
            if gent_form:
                result_tokens.append(gent_form.word)
            else:
                result_tokens.append(token)
        return " ".join(result_tokens)

    @classmethod
    def guess_gender_pronoun(cls, name: str) -> str:
        if name:
            parsed = morph.parse(name)[0]
            gender = parsed.tag.gender
            if gender == 'masc':
                return 'який'
            elif gender == 'femn':
                return 'яка'
            else:
                return 'що'
        return ""

    def name_crusher(self, full_name: str) -> dict:
        tokens = full_name.split()
        last_name = tokens[0] if len(tokens) > 0 else ""
        first_name = tokens[1] if len(tokens) > 1 else ""
        middle_name = tokens[2] if len(tokens) > 2 else ""

        return {
            "last_name": self.title_or_empty_str(last_name),
            "upper_last_name": self.upper_or_empty_str(last_name),
            "first_name": self.title_or_empty_str(first_name),
            "upper_first_name": self.upper_or_empty_str(first_name),
            "middle_name": self.title_or_empty_str(middle_name),
            "upper_middle_name": self.upper_or_empty_str(middle_name),
            "upper_director_name": self.upper_or_empty_str(full_name),
            "gender_pronoun": self.guess_gender_pronoun(first_name),
            "genitive_director_name": self.to_genitive(full_name),
            "upper_genitive_director_name": self.upper_or_empty_str(self.to_genitive(full_name)),
        }

    def upper_company_name(self, company_name: str) -> dict:
        return {
            "upper_company_name": self.upper_or_empty_str(company_name),
        }

    def date_dict_generator(self) -> dict:
        return {
            "start_day": f"{self.start_date.day:02d}",
            "start_month": self.MONTH_NAMES_UA[self.start_date.month],
            "start_year": str(self.start_date.year),
            "expire_day": f"{self.expire_date.day:02d}",
            "expire_month": self.MONTH_NAMES_UA[self.expire_date.month],
            "expire_year": str(self.expire_date.year)
        }

    def document_number_generator(self, business_entity_id) -> dict:
        return {
            "document_number": f"{self.start_date.day:02d}/{self.start_date.month}{business_entity_id}",
        }

    def replacements_generator(self) -> dict:

        bank_dict = self._model_to_dict_or_empty(
            self.bank,
            fields=["name", "mfo"]
        )
        business_entity_dict = self._model_to_dict_or_empty(
            self.business_entity,
            fields=[
                "business_entity",
                "edrpou",
                "company_name",
                "director_name",
                "address",
                "email",
                "phone",
                "iban",
            ]
        )

        name_crasher_dict = self.name_crusher(business_entity_dict.get("director_name"))
        upper_company_name_dict = self.upper_company_name(business_entity_dict.get("company_name"))
        date_dict = self.date_dict_generator()
        document_number_dict = self.document_number_generator(self.business_entity.id)

        merged_dict = {
            **bank_dict,
            **business_entity_dict,
            **name_crasher_dict,
            **upper_company_name_dict,
            **date_dict,
            **document_number_dict
        }
        return merged_dict

    def cars_info_generator(self) -> list:
        vehicles_list = self.vehicles.values_list(
            "brand", "model", "number", "year"
        )
        return list(vehicles_list)


class WordDocManager:
    def __init__(self, template_path):
        self.doc = Document(template_path)

    def replace_placeholders(self, replacement: dict):
        doc = self.doc
        docx_replace(
            doc=doc,
            **replacement,
        )

    @classmethod
    def find_table_index_by_header(cls, doc) -> int:
        search_words = ["Марка, модель", "Реєстраційний номер"]
        for table_index, table in enumerate(doc.tables):
            table_text = " ".join(cell.text for row in table.rows for cell in row.cells)
            if all(word in table_text for word in search_words):
                return table_index
        return -1

    def add_car_info_to_table(self, cars_info: list):
        doc = self.doc

        table_index = self.find_table_index_by_header(doc)

        table = doc.tables[table_index]

        for count, car in enumerate(cars_info, start=1):
            row_cells = table.add_row().cells
            row_cells[0].paragraphs[0].add_run(str(count)).bold = True
            row_cells[1].paragraphs[0].add_run(str(car[0]) + ", " + str(car[1])).bold = True
            row_cells[2].paragraphs[0].add_run(str(car[2])).bold = True
            row_cells[3].paragraphs[0].add_run(str(car[3])).bold = True

    @staticmethod
    def create_output_filename(template: Templates, business_entity: BusinessEntities, contract_index) -> str:
        if business_entity.company_name:
            return f"{template.name}_{business_entity.company_name}_{contract_index}.docx"

        return f"{template.name}_{business_entity.director_name}_{contract_index}.docx"

    def save_word_file(self, filename: str = "output.docx") -> str:
        doc = self.doc
        output_path = os.path.join(settings.MEDIA_ROOT, "documents", filename)

        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        doc.save(output_path)
        return output_path


class ContractTemplatesSearchFormView(BusinessEntityMixin, ListView):
    model = Templates
    template_name = 'contract_templates/partials/search_form.html'
    context_object_name = 'templates'
    paginate_by = 5

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['business_entity'] = self.business_entity
        return context


class ContractTemplatesSearchView(BusinessEntityMixin, ListView):
    model = Templates
    template_name = 'contract_templates/partials/search_list.html'
    context_object_name = 'templates'
    paginate_by = 5

    def get_queryset(self):
        query = self.request.GET.get('searchContracts', '')
        return Templates.objects.filter(Q(name__icontains=query))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['query'] = self.request.GET.get('searchContracts', '')
        context['business_entity'] = self.business_entity
        return context


class ContractDocumentDetailListView(BusinessEntityMixin, TemplateView):
    template_name = 'documents/detail_list.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        contracts = Documents.objects.filter(
            contract__business_entities=self.business_entity
        )
        context.update({
            'business_entity': self.business_entity,
            'contracts': contracts,
        })
        return context


class ContractCreateView(BusinessEntityMixin, VehicleMixin, TemplateMixin, TemplateView):
    template_name = "contracts/partials/create_form.html"

    def get(self, request, *args, **kwargs):
        context = self.get_context_data(**kwargs)
        context['time_range_form'] = ContractTimeRangeForm()
        context['business_entity'] = self.business_entity
        context['template'] = self.template_obj
        return self.render_to_response(context)

    def post(self, request, *args, **kwargs):
        time_range_form = ContractTimeRangeForm(request.POST)
        if time_range_form.is_valid():
            vehicles_with_entities = self.vehicles_with_business_entity(business_entity=self.business_entity)

            start_date = time_range_form.cleaned_data["start_date"]
            expire_date = time_range_form.cleaned_data["end_date"]

            contract = Contracts.objects.create(
                business_entities=self.business_entity,
                template=self.template_obj,
                start_date=start_date,
                end_date=expire_date,
            )
            contract.save()

            replacement_manager = ReplacementManager(
                start_date=start_date,
                expire_date=expire_date,
                business_entity=self.business_entity,
                bank=self.business_entity.bank,
                vehicles=vehicles_with_entities,
            )
            word_manager = WordDocManager(template_path=self.template_obj.path.path)

            replacement_dict = replacement_manager.replacements_generator()
            cars_dict = replacement_manager.cars_info_generator()

            word_manager.replace_placeholders(replacement_dict)
            word_manager.add_car_info_to_table(cars_dict)

            filename = word_manager.create_output_filename(
                self.template_obj,
                self.business_entity,
                contract.id
            )
            output = word_manager.save_word_file(filename=filename)

            document_entity = Documents.objects.create(
                name=filename,
                path=output,
                contract=contract
            )
            document_entity.save()
            documents = Documents.objects.filter(contract__business_entities=self.business_entity)

            context = {
                "time_range_form": time_range_form,
                "business_entity": self.business_entity,
                "template": self.template_obj,
                "documents": documents,
            }
            return render(request, "documents/detail_list.html", context)

        context = self.get_context_data(**kwargs)
        context['time_range_form'] = time_range_form
        return self.render_to_response(context)


class ContractTemplatesListView(ListView):
    model = Templates
    template_name = 'contract_templates/list.html'
    context_object_name = 'templates'
    paginate_by = 8


class TemplateCreateView(CreateView):
    model = Templates
    form_class = TemplatesForm
    template_name = 'contract_templates/create.html'
    success_url = reverse_lazy('contracts:templates')
    pk_url_kwarg = 'template_id'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['template_form'] = context['form']
        context['template'] = self.object
        return context


class TemplateUpdateView(UpdateView):
    model = Templates
    form_class = TemplatesForm
    template_name = 'contract_templates/update.html'
    success_url = reverse_lazy('contracts:templates')
    pk_url_kwarg = 'template_id'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['template_form'] = context['form']
        context['template'] = self.object
        return context


class TemplatesListRedirectView(View):
    def get(self, request, *args, **kwargs):
        response = HttpResponse("")
        redirect_url = reverse("contracts:templates")
        response["HX-Redirect"] = redirect_url
        return response


class TemplatesDownloadView(View):
    def get(self, request, pk):
        template = Templates.objects.get(pk=pk)
        if not template.path:
            raise Http404("File not found.")
        file_path = template.path.path
        if os.path.exists(file_path):
            return FileResponse(
                open(file_path, 'rb'),
                as_attachment=True,
                filename=os.path.basename(file_path)
            )
        else:
            raise Http404("File not found.")
