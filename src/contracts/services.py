import os
import re
import pymorphy3

from django.forms import model_to_dict

from business_entities.models import BusinessEntities
from config import settings
from contracts.models import Templates

morph = pymorphy3.MorphAnalyzer(lang='uk')

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from python_docx_replace import docx_replace

class ReplacementTextHandler:
    MONTH_NAMES_UA = {
        1: "січня",
        2: "лютого",
        3: "березня",
        4: "квітня",
        5: "травня",
        6: "червня",
        7: "липня",
        8: "серпня",
        9: "вересня",
        10: "жовтня",
        11: "листопада",
        12: "грудня",
    }

    @classmethod
    def _model_to_dict_or_empty(cls, instance, fields, default="") -> dict:
        if instance is not None:
            return model_to_dict(instance, fields=fields)
        return {field: default for field in fields}

    @classmethod
    def upper_or_empty_str(cls, text: str) -> str:
        return text.upper() if text else ""

    @classmethod
    def title_or_empty_str(cls, text: str) -> str:
        return text.title() if text else ""

    @classmethod
    def to_genitive(cls, full_name: str) -> str:
        tokens = full_name.split()
        result_tokens = []

        for token in tokens:
            parsed = morph.parse(token)
            best_parse = parsed[0]
            gent_form = best_parse.inflect({'gent'})
            if gent_form:
                result_tokens.append(gent_form.word)
            else:
                result_tokens.append(token)
        return " ".join(result_tokens)

    @classmethod
    def guess_gender_pronoun(cls, name: str) -> str:
        if name:
            parsed = morph.parse(name)[0]
            gender = parsed.tag.gender
            if gender == 'masc':
                return 'який'
            elif gender == 'femn':
                return 'яка'
            else:
                return 'що'
        return ""


class ReplacementManager(ReplacementTextHandler):

    def __init__(
            self, start_date, expire_date, business_entity, vehicles, bank
    ):
        self.start_date = start_date
        self.expire_date = expire_date
        self.bank = bank
        self.business_entity = business_entity
        self.vehicles = vehicles

    def format_phone_text(self) -> str:
        phones = re.split(r"[,\s]+", self.business_entity.phone.strip())
        results = []

        for phone in phones:
            if not phone:
                continue
            royal_phone = self.format_phone_number(phone)
            if royal_phone:
                results.append(royal_phone)
            else:
                results.append(phone)

        return ", ".join(results)

    def name_crusher(self, full_name: str) -> dict:
        tokens = full_name.split()
        last_name = tokens[0] if len(tokens) > 0 else ""
        first_name = tokens[1] if len(tokens) > 1 else ""
        middle_name = tokens[2] if len(tokens) > 2 else ""

        return {
            "last_name": self.title_or_empty_str(last_name),
            "upper_last_name": self.upper_or_empty_str(last_name),
            "first_name": self.title_or_empty_str(first_name),
            "upper_first_name": self.upper_or_empty_str(first_name),
            "middle_name": self.title_or_empty_str(middle_name),
            "upper_middle_name": self.upper_or_empty_str(middle_name),
            "upper_director_name": self.upper_or_empty_str(full_name),
            "gender_pronoun": self.guess_gender_pronoun(first_name),
            "genitive_director_name": self.to_genitive(full_name),
            "upper_genitive_director_name": self.upper_or_empty_str(self.to_genitive(full_name)),
        }

    def upper_company_name(self, company_name: str) -> dict:
        return {
            "upper_company_name": self.upper_or_empty_str(company_name),
        }

    def date_dict_generator(self) -> dict:
        return {
            "start_day": f"{self.start_date.day:02d}",
            "start_month": self.MONTH_NAMES_UA[self.start_date.month],
            "start_year": str(self.start_date.year),
            "expire_day": f"{self.expire_date.day:02d}",
            "expire_month": self.MONTH_NAMES_UA[self.expire_date.month],
            "expire_year": str(self.expire_date.year)
        }

    def document_number_generator(self, business_entity_id) -> dict:
        return {
            "document_number": f"{self.start_date.day:02d}/{self.start_date.month}{business_entity_id}",
        }

    def get_business_entity_dict(self) -> dict:
        business_entity_dict = self._model_to_dict_or_empty(
            self.business_entity,
            fields=[
                "business_entity",
                "edrpou",
                "company_name",
                "director_name",
                "address",
                "email",
                "phone",
                "iban",
            ]
        )
        return business_entity_dict

    def get_bank_dict(self) -> dict:
        bank_dict = self._model_to_dict_or_empty(
            self.bank,
            fields=["name", "mfo"]
        )
        bank_dict["bank_name"] = bank_dict.pop("name")
        return bank_dict

    def replacements_generator(self) -> dict:
        bank_dict = self.get_bank_dict()
        business_entity_dict = self.get_business_entity_dict()
        name_crasher_dict = self.name_crusher(business_entity_dict.get("director_name"))
        upper_company_name_dict = self.upper_company_name(business_entity_dict.get("company_name"))
        date_dict = self.date_dict_generator()
        document_number_dict = self.document_number_generator(self.business_entity.id)

        merged_dict = {
            **bank_dict,
            **business_entity_dict,
            **name_crasher_dict,
            **upper_company_name_dict,
            **date_dict,
            **document_number_dict
        }
        return merged_dict

    def cars_info_generator(self) -> list:
        vehicles_list = self.vehicles.values_list(
            "brand", "model", "number", "year"
        )
        return list(vehicles_list)


class RoyalReplacement(ReplacementManager):

    @classmethod
    def format_phone_number(cls, phone):
        digits = "".join(ch for ch in phone if ch.isdigit())

        if digits.startswith("380") and len(digits) >= 11:
            digits = digits[2:]
        if len(digits) == 10 and digits.startswith("0"):
            operator_code = digits[1:4]

            local_7_digits = digits[4:]

            if len(local_7_digits) == 7:
                part_abc = local_7_digits[:3]
                part_de = local_7_digits[3:5]
                part_fg = local_7_digits[5:]

                royal_phone = f"38 ({operator_code}) {part_abc} {part_de} {part_fg}"

                return royal_phone

        return ""

    def get_business_entity_dict(self) -> dict:
        business_entity_dict = {
            "business_entity": self.business_entity.business_entity if self.business_entity.business_entity else "",
            "edrpou": f"Код ЄДРПОУ {self.business_entity.edrpou}\n" if self.business_entity.edrpou else "",
            "company_name": self.business_entity.company_name if self.business_entity.company_name else "",
            "director_name": self.business_entity.director_name if self.business_entity.director_name else "",
            "address": f"Адреса: {self.business_entity.address}\n" if self.business_entity.address else "",
            "email": f"{self.business_entity.email}\n" if self.business_entity.email else "",
            "phone": f"Тел. {self.format_phone_text()}\n" if self.business_entity.phone else "",
            "iban": f"IBAN {self.business_entity.iban}\n" if self.business_entity.iban else "",
        }
        return business_entity_dict

    def get_bank_dict(self) -> dict:
        if self.bank:
            bank_dict = {
                "bank_name": f"в{self.bank.name};" if self.bank.name else "",
                "mfo": f" МФО {self.bank.mfo}\n" if self.bank.mfo else "",
            }
            return bank_dict
        return {"bank_name": "",
                "mfo": ""}


class RolandReplacement(ReplacementManager):

    @classmethod
    def format_phone_number(cls, phone):
        digits = "".join(ch for ch in phone if ch.isdigit())

        if digits.startswith("380") and len(digits) >= 11:
            digits = digits[2:]
        if len(digits) == 10 and digits.startswith("0"):
            operator_code = digits[1:4]

            local_7_digits = digits[4:]

            if len(local_7_digits) == 7:
                part_abc = local_7_digits[:3]
                part_d = local_7_digits[3:4]
                part_efg = local_7_digits[4:]

                roland_phone = f"+38 ({operator_code}) {part_abc}-{part_d}-{part_efg}"

                return roland_phone
        return ""

    def get_business_entity_dict(self) -> dict:
        business_entity_dict = {
            "business_entity": self.business_entity.business_entity if self.business_entity.business_entity else "",
            "edrpou": f"Код ЄДРПОУ {self.business_entity.edrpou}\n" if self.business_entity.edrpou else "",
            "company_name": self.business_entity.company_name if self.business_entity.company_name else "",
            "director_name": self.business_entity.director_name if self.business_entity.director_name else "",
            "address": f"{self.business_entity.address}\n" if self.business_entity.address else "",
            "email": f"Е-mail: {self.business_entity.email}\n" if self.business_entity.email else "",
            "phone": f"Тел. {self.format_phone_text()}\n" if self.business_entity.phone else "",
            "iban": f"П/р  {self.business_entity.iban}\n" if self.business_entity.iban else "",
        }
        return business_entity_dict

    def get_bank_dict(self) -> dict:
        if self.bank:
            bank_dict = {
                "bank_name": f"в{self.bank.name}," if self.bank.name else "",
                "mfo": f" МФО {self.bank.mfo}\n" if self.bank.mfo else "",
            }
            return bank_dict
        return {"bank_name": "",
                "mfo": ""}


class WordDocManager:
    def __init__(self, template_path):
        self.doc = Document(template_path)

    def replace_placeholders(self, replacement: dict):
        doc = self.doc
        docx_replace(
            doc=doc,
            **replacement,
        )

    def find_font_style(self):
        doc = self.doc
        if doc.paragraphs:
            first_paragraph = doc.paragraphs[0]
            for run in first_paragraph.runs:
                if len(run.text) > 0:
                    font_name = run.font.name

                    if font_name is None and run.style:
                        font_name = run.style.font.name

                    if font_name is None and first_paragraph.style:
                        font_name = first_paragraph.style.font.name

                    return font_name

    @classmethod
    def find_table_index_by_header(cls, doc) -> int:
        search_words = ["Марка, модель", "Реєстраційний номер"]
        for table_index, table in enumerate(doc.tables):
            table_text = " ".join(cell.text for row in table.rows for cell in row.cells)
            if all(word in table_text for word in search_words):
                return table_index
        return -1

    def add_car_info_to_table(self, cars_info: list):
        doc = self.doc

        table_index = self.find_table_index_by_header(doc)

        table = doc.tables[table_index]

        for count, car in enumerate(cars_info, start=1):
            row_cells = table.add_row().cells
            row_cells[0].paragraphs[0].add_run(str(count)).bold = True
            row_cells[1].paragraphs[0].add_run(str(car[0]) + ", " + str(car[1])).bold = True
            row_cells[2].paragraphs[0].add_run(str(car[2])).bold = True
            row_cells[3].paragraphs[0].add_run(str(car[3])).bold = True

        font_name = self.find_font_style()

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.font.name = font_name

    @staticmethod
    def create_output_filename(template: Templates, business_entity: BusinessEntities, contract_index) -> str:
        if business_entity.company_name:
            return f"{template.name}_{business_entity.company_name}_{contract_index}.docx"

        return f"{template.name}_{business_entity.director_name}_{contract_index}.docx"

    def save_word_file(self, filename: str = "output.docx") -> str:
        doc = self.doc
        output_path = os.path.join(settings.MEDIA_ROOT, "documents", filename)

        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        doc.save(output_path)
        return output_path
